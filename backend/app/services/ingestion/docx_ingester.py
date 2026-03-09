from __future__ import annotations

"""DOCX document ingester using python-docx.

This ingester is responsible for:
    - Reading the encrypted DOCX bytes from disk.
    - Decrypting them in memory using the shared AES-256-GCM utilities.
    - Extracting plain text from document paragraphs.
    - Returning an :class:`IngestionResult` with the concatenated text and
      lightweight, non-PII metadata (e.g., paragraph count).
"""

import asyncio
from io import BytesIO
from pathlib import Path
from typing import Any, Dict

from docx import Document as DocxDocument  # type: ignore[import]

from ...utils.crypto import decrypt
from .base import BaseIngester, IngestionResult


class DocxIngester(BaseIngester):
    """Ingests encrypted DOCX documents and extracts their textual content."""

    async def ingest(
        self,
        file_path: Path,
        *,
        mime_type: str,
        file_format: str,
    ) -> IngestionResult:
        """Ingest the DOCX at ``file_path`` and return extracted content."""

        return await asyncio.to_thread(
            self._ingest_sync,
            file_path,
            mime_type,
            file_format,
        )

    def _ingest_sync(
        self,
        file_path: Path,
        mime_type: str,
        file_format: str,
    ) -> IngestionResult:
        """Synchronous part of DOCX ingestion, run in a worker thread."""

        encrypted_bytes = file_path.read_bytes()
        docx_bytes = decrypt(encrypted_bytes)

        document = DocxDocument(BytesIO(docx_bytes))

        paragraphs = [p.text or "" for p in document.paragraphs]
        full_text = "\n".join(paragraphs)

        metadata: Dict[str, Any] = {
            "paragraph_count": len(document.paragraphs),
            "mime_type": mime_type,
            "file_format": file_format,
        }

        return IngestionResult(text=full_text, metadata=metadata)

