from __future__ import annotations

"""
Lightweight file readers for the ``scan_file`` MCP tool.

This module intentionally keeps the dependency surface small: only
pure-Python parsers are used (``pypdf``, ``python-docx``) plus the
standard library. Heavy extractors (OCR via PaddleOCR/Tesseract, audio
via Whisper, video, scanned PDF pages) are out of scope for the MCP
server — they belong in the full septum-api pipeline where the cost
of loading multi-gigabyte models is acceptable.

Content-based MIME detection is *not* used here because it would pull
in ``python-magic`` (libmagic system dependency). The MCP tool is
invoked from editors that already know the file's extension, so
extension-based dispatch is enough for this entrypoint.

All readers return :class:`FileReadResult` objects so the caller can
distinguish "format unsupported" from "file missing" from "file was
read successfully but is empty".
"""

import csv
import io
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Dict

SUPPORTED_EXTENSIONS: tuple[str, ...] = (
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".json",
    ".pdf",
    ".docx",
)


@dataclass
class FileReadResult:
    """Outcome of a :func:`read_file` call.

    ``ok=False`` means the file could not be turned into text. Callers
    should surface ``error`` verbatim to the MCP client so the user
    can see the failure reason in their editor.
    """

    ok: bool
    text: str = ""
    format: str = ""
    error: str = ""


def read_file(path: str | Path) -> FileReadResult:
    """Read ``path`` and return its textual content.

    Format dispatch is based on the file extension. Unknown extensions
    return an ``ok=False`` result listing the supported formats so the
    caller can show a precise error message.
    """
    file_path = Path(path)
    if not file_path.exists():
        return FileReadResult(ok=False, error=f"File not found: {file_path}")
    if not file_path.is_file():
        return FileReadResult(ok=False, error=f"Not a regular file: {file_path}")

    extension = file_path.suffix.lower()
    reader = _EXTENSION_READERS.get(extension)
    if reader is None:
        supported = ", ".join(SUPPORTED_EXTENSIONS)
        return FileReadResult(
            ok=False,
            error=(
                f"Unsupported file extension '{extension}'. "
                f"septum-mcp supports: {supported}"
            ),
        )

    try:
        text = reader(file_path)
    except Exception as exc:
        return FileReadResult(
            ok=False,
            format=extension.lstrip("."),
            error=f"Failed to read {extension} file: {exc}",
        )

    return FileReadResult(ok=True, text=text, format=extension.lstrip("."))


def _read_plain_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read_csv(path: Path) -> str:
    """Return CSV content as tab-joined lines for detector-friendly input."""
    lines: list[str] = []
    with path.open("r", encoding="utf-8", errors="replace", newline="") as handle:
        reader = csv.reader(handle)
        for row in reader:
            lines.append("\t".join(row))
    return "\n".join(lines)


def _read_json(path: Path) -> str:
    """Pretty-print JSON so string values remain detectable by the pipeline."""
    raw = path.read_text(encoding="utf-8", errors="replace")
    try:
        parsed = json.loads(raw)
    except json.JSONDecodeError:
        return raw
    buffer = io.StringIO()
    json.dump(parsed, buffer, ensure_ascii=False, indent=2)
    return buffer.getvalue()


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n".join(pages)


def _read_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts: list[str] = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells if cell.text]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


_EXTENSION_READERS: Dict[str, Callable[[Path], str]] = {
    ".txt": _read_plain_text,
    ".md": _read_plain_text,
    ".markdown": _read_plain_text,
    ".csv": _read_csv,
    ".json": _read_json,
    ".pdf": _read_pdf,
    ".docx": _read_docx,
}
